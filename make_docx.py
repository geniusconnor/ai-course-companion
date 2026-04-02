from docx import Document
from docx.shared import Pt, Inches

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

WD_ALIGN_PARAGRAPH_CENTER = 1

def add_para(text, bold=False, center=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH_CENTER
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = bold
    return p

# Title Page
add_para('')
add_para('')
add_para('')
add_para('Overcoming the Productivity Challenge in Product Life-Extending Operations:', bold=True, center=True)
add_para('A Critical Analysis', bold=True, center=True)
add_para('')
add_para('')
add_para('Student Name', center=True)
add_para('BAN-5003 Operations Analytics', center=True)
add_para('Trine University', center=True)
add_para('March 29, 2026', center=True)

doc.add_page_break()

SECTIONS = [
    ('Introduction', [
        'Every company, and every product it makes, follows a life cycle. Sales climb, peak, and eventually fall \u2014 the familiar curve that shows up in every business school textbook. Most executives know this shape intuitively. What they rarely examine is what happens inside the mountain: the operational layers buried beneath the peak that determine whether a product creates value once, or keeps creating it long after the first sale.',
        'Product life-extending operations (PLEO) \u2014 remanufacturing, repair, refurbishing, reconditioning, overhaul, and maintenance \u2014 live in that interior. They are not the headline. They do not show up on the product launch slide deck. Because they are overlooked, the ROI potential they carry for the business tends to be underestimated. Most CEOs put their attention on sales conversion: it is immediate, visible, and tied directly to quarterly performance. That focus is rational in the short run. But for organizations operating with a long-term orientation \u2014 and for an economy trying to reduce waste at scale \u2014 PLEO and the circular economy logic behind it deserve a seat at the table.',
        'Gr\u00f3f and Netland (2025) make the case for that seat rigorously. Their central argument is that PLEO underperforms original manufacturing not because the problem is inherently unsolvable, but because the operations are poorly designed. Their study builds a process framework, develops a theoretical model, and validates both against fourteen European companies. This paper works through that argument and ends with an evaluation of where the model holds, where it leaves a gap, and what organizations can actually do with it.',
    ]),
    ('Key Issues', [
        'The gap Gr\u00f3f and Netland (2025) are trying to explain is real and persistent. PLEO consistently underperform original manufacturing on throughput time, cycle time variability, inventory levels, and resource efficiency. The circular economy case for PLEO is straightforward: industry estimates suggest fewer than 2% of products with reuse potential in Europe are actually reused (European Remanufacturing Network, 2015, as cited in Gr\u00f3f & Netland, 2025). The barrier is economic. PLEO cannot scale if they bleed money relative to linear alternatives.',
        'The standard explanation in the literature is product-induced variability. Products return from use in unpredictable condition, at unpredictable times, in unpredictable quantities (Guide, 2000; Goltsos et al., 2019, as cited in Gr\u00f3f & Netland, 2025). This makes PLEO look like a fundamentally hard problem: the inputs are uncontrollable, so the process is condemned to inefficiency. Most scholarly work in this space has tried to improve planning and decision-making within that constraint \u2014 accepting it as given.',
        'Gr\u00f3f and Netland (2025) do not accept it. They identify a second source of variability, largely absent from the existing literature: the customer. In PLEO focused on asset recovery, where a specific customer product is being restored rather than replaced, the customer participates actively throughout the process. Customers define disassembly scope, approve cost estimates after inspection, supply spare parts, and sit on information requests while operations wait. Each of these is a source of delay that has nothing to do with what condition the product arrived in.',
        'More importantly, the authors argue that both sources of variability are at least partly determined by process design decisions made before any individual product arrives. PLEO productivity is not a ceiling set by reverse logistics. It is partly a choice.',
    ]),
    ('Process Framework', [
        'To give that claim operational content, Gr\u00f3f and Netland (2025) develop a process framework drawn from their nine case companies. It describes 11 activities and the decisions each one involves.',
        'The process starts with core stock \u2014 the decision of whether to buffer incoming products before processing begins. This decision has historically been driven by experience rather than data. In earlier business environments, inventory calls were made by operations managers relying on intuition, or delegated to whoever had the closest contact with first-hand information: a warehouse supervisor reading demand patterns, a sales rep sensing when product would start coming back. That approach was imprecise but workable at small scale. In the post-digitalization era, real-time inventory management systems \u2014 and increasingly, AI-driven demand and return forecasting \u2014 make it possible to align purchasing, sales, and warehouse decisions dynamically rather than reactively. The core stock question is no longer purely a judgment call. It is increasingly a data problem. Asset ownership operations tend to hold large core stocks to guarantee product availability; asset recovery operations typically hold none, because the process cannot begin until the customer\u2019s specific product arrives. That dependency is a form of customer interaction built into the process from the first step.',
        'Incoming inspection and testing follows. Disassembly is next, and here the framework draws a distinction that matters. There are four types: complete, event-specific, condition-based, and customised. These represent a spectrum between two competing objectives. Complete disassembly is the most thorough and the most expensive. At the other end, sampling-based approaches minimize cost by inspecting only the components most likely to fail. The tension is a probability question: what is the acceptable failure rate, and who bears the cost when something is missed? That is a strategic decision for management, not a technical one.',
        'After disassembly comes cleaning, component inspection, and technical evaluation. In these large enterprises, there is at least some opportunity for staff to listen to the product directly. But heavily regulated industries require specialized internal departments or third-party certification bodies, and that makes accelerating the inspection stage genuinely difficult. Certification processes move at their own pace \u2014 regulatory bodies are not known for speed, and in some markets, navigating approvals introduces delays that no amount of internal process optimization can fully absorb.',
        'The customer decision step is unavoidable \u2014 it sits in the critical path of every asset recovery operation. The real question is not whether to include it but how to help customers decide faster and interpret inspection results more clearly. That answer varies sharply by industry, product type, and customer sophistication. You are accountable for the customer\u2019s decision, but you cannot shortcut it without assuming liability you did not intend to carry.',
        'Part-level activities then split into refurbishment or replacement, with sourcing either company-procured or customer-supplied. The practical ideal is straightforward: the right part, fit for purpose, delivered into the customer\u2019s active workflow at minimum cost. The more this process can be quantified, the faster operations run. With AI-driven diagnostics and increasing deployment of physical sensors on the shop floor, this is arguably the part of the PLEO chain most amenable to near-term acceleration.',
        'Across all eleven activities, the common thread is this: modern operational efficiency comes from combining digitalization with precision management. Layering in tools \u2014 AI platforms, SaaS-based workflow systems, real-time sensor data \u2014 optimizes decision nodes for management and improves the experience for the customer. Done well, that is not a trade-off. It is a compounding advantage for the whole operation.',
    ]),
    ('Research Methodology', [
        'Gr\u00f3f and Netland (2025) used a multiple-case study approach based on Eisenhardt (1989), sampling nine European companies with longstanding, profitable PLEO operations for in-depth analysis. The data came from 59 informants across these companies, totaling approximately 60 hours and 47 minutes of structured interviews, ranging from shop floor supervisors to senior management. The research team also collected organizational charts, process maps, facility layouts, annual reports, and technical documentation, and conducted one-day on-site visits at each company.',
        'Analysis used MAXQDA for open coding, starting inductively to systematically review all interview material. Despite sharing many operational similarities, the companies showed substantial variation in actual management processes \u2014 which is what made the study interesting. A post-hoc validation phase brought in five additional companies \u2014 MachineryCo, AeroComponentCo, RailCo, AeroMaintenanceCo, and NationalAirlineCo \u2014 to stress-test the emerging model. None contradicted it.',
    ]),
    ('Critical Evaluation', [
        'The theoretical model Gr\u00f3f and Netland (2025) propose carries real limitations. The sample \u2014 fourteen large, established European OEMs with significant organizational scale \u2014 introduces an inherent industry bias. I can follow the high-level cost-optimization logic; at that level, the direction is sound. But pushing it down into second- and third-tier applications, and expecting the same methodology to hold for smaller organizations, is where the model starts to break down. Here is what I mean.',
        'The first issue is scope. These are organizations where separating a customer decision step from everything else is already a meaningful organizational act. Dedicated departments exist. SOPs are documented. Process design is a real variable. For that context, the framework is genuinely useful.',
        'It does not transfer cleanly to smaller operations, though. I have seen this up close. When a company runs with fewer than ten people, there is no discrete customer decision step to move upstream or eliminate. There is one person handling sales, returns, customer complaints, and logistics \u2014 often in the same afternoon. In that setting, the binding constraint is not process variability. It is cash flow. It is survival. Putting this model into practice too early adds overhead for a team that needs to stay loose. The framework works for companies that have already solved the staying-alive problem. It was not built for companies still working on it.',
        'The second issue is how the model treats customer interaction. It frames customer involvement primarily as a source of process delay. That is partly right. But the more important question is not how much interaction there is. It is which interactions are worth managing at all.',
        'Not all customer feedback carries the same cost or signal. At Ulanzi, we handle returns and defective camera accessories at real volume. We learned this slowly: a portion of complaint-driven customers generates a disproportionate share of service overhead relative to their actual revenue contribution. Some complaints reflect genuine product problems. Many do not \u2014 the product works fine, the customer is dissatisfied for reasons unrelated to the product, and the cost of managing the interaction exceeds anything we could realistically recover from retaining them. A founder running PLEO needs to calculate the actual cost of each customer interaction type against the business value of the customer generating it. The model does not help with that.',
        'The third issue is the circularity trade-off. I grew up in China. I know what it costs to manufacture a consumer product in Shenzhen. For a lot of the goods that move through Western retail, the unit cost is low enough that the reverse logistics chain in a developed country costs more than the product itself. A camera accessory made in Shenzhen, sold through Amazon in the US, returned by a customer in Texas: the shipping cost alone may exceed the product\u2019s landed cost. A warehouse worker in Memphis earns more per hour than the product is worth. In that math, the rational decision is to tell the customer to keep the defective unit and ship a replacement. Several large e-commerce platforms have already done exactly this \u2014 not because they stopped caring about waste, but because the alternative costs more.',
        'The circular economy\u2019s aspirations are worth taking seriously. But economic viability has to come first. A company that loses money pursuing circularity will eventually stop pursuing it. In a meaningful share of consumer goods markets, recovery is a math problem with a clear answer \u2014 and the answer is not always recovery.',
        'Three recommendations follow. First, digital product passports reduce product interaction overhead without changing technical scope \u2014 aviation\u2019s airworthiness documentation is the precedent. Second, pre-arrival customer decision agreements remove mid-process waiting: agree on repair thresholds before the product arrives, not after disassembly. Third, track product interaction intensity as a leading KPI, not just throughput time.',
    ]),
    ('Conclusion', [
        'Gr\u00f3f and Netland (2025) shift the PLEO productivity conversation from what operators receive to what they decide. The process framework is specific enough to apply to real operations. The theoretical model explains why different configurations produce different results.',
        'The model works best for large, established operations where process design is the binding constraint. It has less to offer companies that have not yet reached that scale. It also does not resolve what practitioners actually face: which customer interactions are worth the cost, and whether circular recovery is economically viable at their specific price point. Those questions need a different kind of analysis \u2014 one that starts from the economics of a particular business.',
    ]),
]

for heading, paras in SECTIONS:
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH_CENTER
    h.paragraph_format.line_spacing = Pt(24)
    h.paragraph_format.space_after = Pt(0)
    h.paragraph_format.space_before = Pt(0)
    run = h.add_run(heading)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    for para in paras:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(24)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.first_line_indent = Inches(0.5)
        run = p.add_run(para)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# References
doc.add_page_break()
ref_h = doc.add_paragraph()
ref_h.alignment = WD_ALIGN_PARAGRAPH_CENTER
ref_h.paragraph_format.line_spacing = Pt(24)
ref_h.paragraph_format.space_after = Pt(0)
run = ref_h.add_run('References')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.font.bold = True

REFS = [
    'Gr\u00f3f, C., & Netland, T. (2025). Overcoming the productivity challenge in product life-extending operations: a multiple-case study of European facilities. International Journal of Production Research, 63(21), 8026\u20138046. https://doi.org/10.1080/00207543.2025.2509170',
    'Schmenner, R. W., & Swink, M. L. (1998). On theory in operations management. Journal of Operations Management, 17(1), 97\u2013113. https://doi.org/10.1016/S0272-6963(98)00028-X',
    'Sampson, S. E., & Froehle, C. M. (2006). Foundations and implications of a proposed unified services theory. Production and Operations Management, 15(2), 329\u2013343. https://doi.org/10.1111/j.1937-5956.2006.tb00248.x',
    'Eisenhardt, K. M. (1989). Building theories from case study research. Academy of Management Review, 14(4), 532\u2013550. https://doi.org/10.2307/258557',
    'Guide, V. D. R., Jr. (2000). Production planning and control for remanufacturing: Industry practice and research needs. Journal of Operations Management, 18(4), 467\u2013483. https://doi.org/10.1016/S0272-6963(99)00034-6',
    'Ellen MacArthur Foundation. (2015). Towards a circular economy: Business rationale for an accelerated transition. Ellen MacArthur Foundation.',
]

for ref in REFS:
    rp = doc.add_paragraph()
    rp.paragraph_format.line_spacing = Pt(24)
    rp.paragraph_format.space_after = Pt(0)
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.left_indent = Inches(0.5)
    rp.paragraph_format.first_line_indent = Inches(-0.5)
    run = rp.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.save('e:/伴读书童/courses/assignments/BAN5003-W3-CasePaper-v5.docx')
print('saved')
